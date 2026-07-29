from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase.pdfmetrics import stringWidth
from reportlab.pdfgen import canvas


PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
SIDE_MARGIN_CM = 0.05
TOP_BOTTOM_MARGIN_CM = 2.2
COLUMN_GAP_CM = 0.5
LABELS_PER_ROW = 2
ROWS_PER_PAGE = 7
LABEL_WIDTH_CM = (PAGE_WIDTH_CM - (SIDE_MARGIN_CM * 2) - COLUMN_GAP_CM) / 2
LABEL_HEIGHT_CM = (PAGE_HEIGHT_CM - (TOP_BOTTOM_MARGIN_CM * 2)) / ROWS_PER_PAGE
ESTABLISHMENT_NAME = "MAMOGRAFIA CIS VERDE"
LOGO_PATH = Path(__file__).resolve().parent.parent / "static" / "assets" / "cis-verde-logo.png"


def _format_label(label: dict) -> list[tuple[str, str]]:
    return [
        ("PACIENTE:", label["patient_name"]),
        ("MUNICÍPIO:", f'{label["city"]}   BAIRRO: {label["district"]}'),
        ("DATA DE NASCIMENTO:", label["birth_date"]),
        ("DATA DA REALIZAÇÃO:", label["exam_date"]),
    ]


def _fit_text(text: str, font: str, max_width: float, base_size: int = 10, min_size: int = 7) -> int:
    size = base_size
    while size > min_size and stringWidth(text, font, size) > max_width:
        size -= 1
    return size


def _draw_pdf_label(pdf: canvas.Canvas, label: dict, x: float, y: float, width: float, height: float) -> None:
    left = x + 0.25 * cm
    top = y + height - 0.28 * cm
    if LOGO_PATH.exists():
        pdf.drawImage(ImageReader(str(LOGO_PATH)), left, top - 0.72 * cm, width=1.75 * cm, height=0.92 * cm, mask="auto")

    pdf.setFont("Helvetica-Bold", 9.2)
    pdf.drawString(left + 2.15 * cm, top - 0.35 * cm, ESTABLISHMENT_NAME)

    baseline = top - 1.12 * cm
    for field_index, (title, value) in enumerate(_format_label(label)):
        line = f"{title} {value}"
        font_size = _fit_text(line, "Helvetica-BoldOblique", width - 0.5 * cm, base_size=8.6, min_size=6.8)
        pdf.setFont("Helvetica-BoldOblique", font_size)
        pdf.drawString(left, baseline - (field_index * 0.43 * cm), line)


def generate_pdf(labels: list[dict], output_path: str | Path) -> Path:
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    pdf = canvas.Canvas(str(output), pagesize=A4)
    page_width, page_height = A4
    left_margin = SIDE_MARGIN_CM * cm
    top_margin = TOP_BOTTOM_MARGIN_CM * cm
    label_width = LABEL_WIDTH_CM * cm
    label_height = LABEL_HEIGHT_CM * cm
    gap = COLUMN_GAP_CM * cm

    for index, label in enumerate(labels):
        slot = index % (LABELS_PER_ROW * ROWS_PER_PAGE)
        row = slot // LABELS_PER_ROW
        col = slot % LABELS_PER_ROW
        if index and slot == 0:
            pdf.showPage()

        x = left_margin + col * (label_width + gap)
        y = page_height - top_margin - ((row + 1) * label_height)
        _draw_pdf_label(pdf, label, x, y, label_width, label_height)

    pdf.save()
    return output


def _set_cell_borderless(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "nil")
        borders.append(element)
    tc_pr.append(borders)


def _set_cell_margins(cell, margin_cm: float = 0.18) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    twips = str(int(margin_cm * 567))
    for edge in ("top", "left", "bottom", "right"):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), twips)
        node.set(qn("w:type"), "dxa")


def _add_docx_label(cell, label: dict) -> None:
    cell.text = ""
    header = cell.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(1)
    header.paragraph_format.line_spacing = 1.0
    if LOGO_PATH.exists():
        logo_run = header.add_run()
        logo_run.add_picture(str(LOGO_PATH), width=Cm(1.75))
        header.add_run("  ")
    title = header.add_run(ESTABLISHMENT_NAME)
    title.bold = True
    title.font.name = "Arial"
    title.font.size = Pt(9.5)

    for title_text, value in _format_label(label):
        paragraph = cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 0.92
        run = paragraph.add_run(f"{title_text} {value}")
        run.bold = True
        run.italic = True
        run.font.name = "Arial"
        run.font.size = Pt(7.8 if len(value) > 35 else 8.3)


def generate_docx(labels: list[dict], output_path: str | Path) -> Path:
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.left_margin = Cm(SIDE_MARGIN_CM)
    section.right_margin = Cm(SIDE_MARGIN_CM)
    section.top_margin = Cm(TOP_BOTTOM_MARGIN_CM)
    section.bottom_margin = Cm(TOP_BOTTOM_MARGIN_CM)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(9)

    for page_start in range(0, len(labels), LABELS_PER_ROW * ROWS_PER_PAGE):
        if page_start:
            doc.add_page_break()
        page_labels = labels[page_start : page_start + LABELS_PER_ROW * ROWS_PER_PAGE]
        table = doc.add_table(rows=ROWS_PER_PAGE, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        widths = [Cm(LABEL_WIDTH_CM), Cm(COLUMN_GAP_CM), Cm(LABEL_WIDTH_CM)]
        for row_index, row in enumerate(table.rows):
            row.height = Cm(LABEL_HEIGHT_CM)
            for col_index, cell in enumerate(row.cells):
                cell.width = widths[col_index]
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                _set_cell_borderless(cell)
                _set_cell_margins(cell)
                if col_index == 1:
                    continue
                label_index = page_start + (row_index * LABELS_PER_ROW) + (0 if col_index == 0 else 1)
                if label_index >= len(labels):
                    continue
                _add_docx_label(cell, labels[label_index])

    doc.save(output)
    return output
